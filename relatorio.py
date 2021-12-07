# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime
import io
import csv
from copy import deepcopy
from tkinter import Tk, filedialog, messagebox

def window_select_file():
    root = Tk()
    root.withdraw()
    root.attributes('-topmost', True)
    file = filedialog.askopenfilename()
    return file


class Solicitacao:

    def __init__(self, solicitacao) -> None:
        self.num_laudo = solicitacao[0]
        self.num_sei = solicitacao[1]
        self.solicitante = solicitacao[2]
        self.sistema = solicitacao[3]
        self.tipo = solicitacao[4]
        self.area = solicitacao[5]
        self.valor_total = solicitacao[6]
        

class Proprietario:

    def __init__(self, proprietario_info) -> None:
        self.nome, self.cpf = proprietario_info


class Imovel:

    def __init__(self, imovel) -> None:
        self.endereco, self.municipio, self.cep = imovel


class Avaliacao:

    def __init__(self, avaliacao) -> None:
        self.descricao_regiao, self.descricao_area, self.desempenho,\
             self.fundamentacao, self.precisao, self.num_variaveis = avaliacao     


class Relatorio:

    def __init__(self, laudo_info, proprietario_info, imovel_info, avaliacao) -> None:
        self.solicitacao = Solicitacao(laudo_info)
        self.proprietario = Proprietario(proprietario_info)
        self.imovel = Imovel(imovel_info)
        self.avaliacao = Avaliacao(avaliacao)
        self.doc = Document()
        paragraph = self.doc.add_paragraph()
        paragraph.style = self.doc.styles['Heading 4']

    def title(self):
        self.doc.add_heading('Laudo GDS ' + str(self.solicitacao.num_laudo), 0 )

    @property
    def seleciona_tipo(self):
        if self.solicitacao.tipo == 0:
            return "servidão administrativa"
        elif self.solicitacao.tipo == 1:
            return "desapropriação"


    def cabecalho(self):

        def processo():
            p = self.doc.add_paragraph('Processo: Solicitado via SEI Nº ' + self.solicitacao.num_sei + '\n')
            p.add_run('Solicitante: ' + self.solicitacao.solicitante + '\n')
            p.add_run('Objeto: Avaliação de Imóvel para ' + self.seleciona_tipo)
            p.add_run('\n')
            p.add_run('Objetivo: Determinação do Justo Valor Indenizatório')

        def proprietario():
            p = self.doc.add_paragraph('Dados do proprietário:' + '\n')
            p.add_run('Proprietário: ' + self.proprietario.nome + '\n')
            p.add_run('CPF: ' + self.proprietario.cpf)

        def endereco():
            p = self.doc.add_paragraph('Endereço completo do imóvel:' + '\n')
            p.add_run('Endereço: : ' + self.imovel.endereco + '\n')
            p.add_run('Município: ' + self.imovel.municipio + '\n')          
            p.add_run('CEP: ' + self.imovel.cep + '\n')          
            p.add_run('Área para ' + self.seleciona_tipo + ': ' + str(self.solicitacao.area) + 'm²') 

        def metodos():
            p = self.doc.add_paragraph('Método(s) utilizado(s):' + '\n')
            p.add_run('Método Comparativo Direto de Dados de Mercado')

        def resultado():
            p = self.doc.add_paragraph('Resultado da avaliação:' + '\n')
            p.add_run('Valor indenizatório: ' + self.solicitacao.valor_total)

        def local_e_data():
            p = self.doc.add_paragraph('Local e data do Laudo de avaliação:' + '\n')
            p.add_run('Recife, ' + datetime.today().strftime('%d/%m/%Y'))

        processo() 
        proprietario()
        endereco()       
        metodos()
        resultado()
        local_e_data()

    def capitulo_01(self):
        self.doc.add_heading('IDENTIFICAÇÃO DO SOLICITANTE', level=1)
        p = self.doc.add_paragraph('Solicitante: ' + self.solicitacao.solicitante + '\n')
        p.add_run('Sistema: ' + self.solicitacao.sistema)

    def capitulo_02(self):
        self.doc.add_heading('FINALIDADE DO LAUDO', level=1)
        p = self.doc.add_paragraph(f"A finalidade deste laudo é avaliar o imóvel para fins de indenização fruto de {self.seleciona_tipo}.")

    def capitulo_03(self):
        self.doc.add_heading('OBJETIVO DA AVALIAÇÃO', level=1)
        lines = io.open("templates\capitulo03.txt", "r", encoding="utf8")
        p = self.doc.add_paragraph(lines.read())
   
    def capitulo_04(self):
        self.doc.add_heading('PRESSUPOSTOS, RESSALVAS E FATORES LIMITANTES', level=1)
        lines = io.open("templates\capitulo04.txt", "r", encoding="utf8")
        for id, line in enumerate(lines):
            if id == 0:
                self.doc.add_paragraph(line)
            else:
                self.doc.add_paragraph(line, style='List Bullet')

    def capitulo_05(self):
        self.doc.add_heading('DESCRIÇÃO DA REGIÃO', level=1)
        p = self.doc.add_paragraph(self.avaliacao.descricao_regiao)
    
    def capitulo_06(self):
        self.doc.add_heading('IDENTIFICAÇÃO E CARACTERIZAÇÃO DO IMÓVEL AVALIANDO', level=1)
        p = self.doc.add_paragraph(self.avaliacao.descricao_area)

    def capitulo_07(self):
        self.doc.add_heading('ANÁLISE DO MERCADO', level=1)
        lines = io.open("templates\capitulo07.txt", "r", encoding="utf8")
        for id, line in enumerate(lines):
            if id == 0:
                self.doc.add_paragraph(line)
            else:
                self.doc.add_paragraph(line + self.avaliacao.desempenho[id-1])

    def tabela_fundamentacao(self):
        p = self.doc.add_paragraph("\n\n Tabela 1 – Grau de fundamentação no caso de utilização de modelos de regressão linear")      
        fundamentacao = self._csv_to_tuples("templates\\fundamentacao.CSV")
        records = []
        for element, pontos in zip(fundamentacao, self.avaliacao.fundamentacao):
            element += ((pontos), )
            records.append(element)
        records.append(('', '', '', '', 'TOTAL', sum(self.avaliacao.fundamentacao)))
        self._create_table(records, ['Item', 'Descrição', 'III', 'II', 'I', 'Pontos'])

    def enquadramento_fundamentacao(self):
            p = self.doc.add_paragraph("\n\n Tabela 2 – Enquadramento do laudo segundo seu grau de fundamentação no caso de modelos de regressão linear")
            
            records = (
                ('Pontos mínimos', '16', '11', '6'),
                ('Itens obrigatórios', 'Itens 2,4,5 e 6 no Grau III e os demais no mínimo no Grau II.', 'Itens 2, 4, 5 e 6 no mínimo no Grau II e os demais no mínimo no Grau I.', 'Todos, no mínimo no grau I')
                )
            
            self._create_table(records, ['Graus', 'III', 'II', 'I'])
        
    @property
    def grau_fundamentacao(self):
        pontos = self.avaliacao.fundamentacao
        if all(i >= j for i, j in zip(pontos, [2, 3, 2, 3, 3, 3])):
            return 'III'
        elif all(i >= j for i, j in zip(pontos, [1, 2, 1, 2, 2, 2])):
            return 'II'
        else:
            return 'I'

    @property
    def grau_precisao(self):
        if self.avaliacao.precisao <= 30:
            return 'III'
        elif  30 < self.precisao <= 40:
            return 'II'
        else:
            return 'I'

    def _create_table(self, records, titulos):

        def center_style(cell):
            cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        table = self.doc.add_table(rows=1, cols=len(records[0]))
        hdr_cells = table.rows[0].cells
        for head, titulo in zip(hdr_cells, titulos):
            head.text = titulo
            center_style(head)
        for record in records:
            row_cells = table.add_row().cells
            for row, cell in zip(row_cells, record):
                row.text = str(cell)
                center_style(row)

    def _csv_to_tuples(self, file):
        with open(file) as f:
            return [tuple(line) for line in csv.reader(f, delimiter=";")]       
            
    def enquadramento_precisao(self):
            p = self.doc.add_paragraph("\n Tabela 5 – Grau de precisão nos casos de utilização de modelos de regressão linear ou do tratamento por fatores")
            
            records = (
                ("Amplitude do intervalo de confiança de 80% em torno da estimativa de tendência central", '≤ 30%', '≤ 40%', '≤ 50%'),
                ('', '', 'Amplitude do intervalo de confiança:', f'{self.avaliacao.precisao}' + '%'),
                ('', '', 'Grau de precisão:', self.grau_precisao),
                )

            self._create_table(records, ['Descrição', 'III', 'II', 'I'])

    def capitulo_08(self):
        self.doc.add_heading('ESPECIFICAÇÃO DA AVALIAÇÃO', level=1)
        lines = io.open("templates\capitulo08.txt", "r", encoding="utf8")
        for id, line in enumerate(lines):
            if id == 0:
                p = self.doc.add_paragraph(line)
            else:
                p.add_run(line)
        
        def fundamentacao(self):
            self.doc.add_heading('Quanto a fundamentação:', level=2)
            p = self.doc.add_paragraph(f"Conforme a NBR 14.653-2, a presente avaliação está enquadrada no Grau {self.grau_fundamentacao} de fundamentação, no caso de utilização do tratamento por inferência estatística, segundo tabelas da Norma transcritas abaixo:")

        def precisao(self):
            self.doc.add_heading('\nQuanto a precisao:', level=2)
            p = self.doc.add_paragraph(f"Conforme a NBR 14.653-2, a presente avaliação está enquadrada no Grau {self.grau_precisao} da estimativa de valor, no caso de utilização do tratamento por fatores, segundo tabela da Norma transcrita abaixo:")

        fundamentacao(self)
        self.tabela_fundamentacao()
        self.enquadramento_fundamentacao()
        precisao(self)
        self.enquadramento_precisao()

    def add_table_from_file(self, indices):
        file = window_select_file()
        tabela_caract_avaliando = Document(file)
        for id in list(indices):
            table = tabela_caract_avaliando.tables[id]
            paragraph = self.doc.add_paragraph()
            paragraph._p.addnext(table._tbl)        

    def capitulo_09(self):
        self.doc.add_heading('AVALIAÇÃO DO IMÓVEL', level=1)
        lines = io.open("templates\capitulo09.txt", "r", encoding="utf8")
        for id, line in enumerate(lines):
            if id == 0:
                self.doc.add_paragraph(line)
            else:
                self.doc.add_paragraph(line)

        p = self.doc.add_paragraph("Descrição das variáveis:")
        start = 17
        conjunto = range(start, start + self.avaliacao.num_variaveis + 1)
        self.add_table_from_file(conjunto)
        
        p = self.doc.add_paragraph("")
        p = self.doc.add_paragraph("Características do Avaliando:")
        self.add_table_from_file([0])

    def save(self):
        self.doc.save("demo2.docx")

    def gerar_relatorio(self):
        self.title()
        self.cabecalho()
        self.capitulo_01()
        self.capitulo_02()
        self.capitulo_03()
        self.capitulo_04()
        self.capitulo_05()
        self.capitulo_06()
        self.capitulo_07()
        self.capitulo_08()
        self.capitulo_09()
        self.save()
                
